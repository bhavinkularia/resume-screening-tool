"""
ATS — Applicant Tracking System
Rule-based resume screening. Skills matched ONLY from SKILL_LIBRARY.
No free-form token extraction. No noise. Clean modular design.
Claude API used ONLY for score refinement + Strengths & Gaps per shortlisted candidate.
Final Score = 70% base (rule-based) + 30% Claude-refined.
"""

import io
import os
import re
from collections import defaultdict

import anthropic

import streamlit as st
import pdfplumber
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===========================================================================
# SKILL LIBRARY
# Single source of truth. Only these exact strings will ever match.
# Add domain-specific skills here to extend coverage.
# ===========================================================================

SKILL_LIBRARY: dict[str, list[str]] = {
    "programming": [
        "python", "r", "java", "javascript", "typescript", "c", "c++", "c#",
        "golang", "go", "rust", "scala", "kotlin", "swift", "php", "ruby",
        "perl", "matlab", "bash", "shell", "powershell", "vba", "sas", "stata",
        "julia", "cobol", "assembly",
    ],
    "web": [
        "html", "css", "react", "react.js", "angular", "vue", "vue.js",
        "next.js", "node.js", "express", "django", "flask", "fastapi",
        "spring", "asp.net", "laravel", "rails", "graphql", "rest api",
        "restful api", "soap", "webpack", "tailwind", "bootstrap", "jquery",
        "redux", "svelte",
    ],
    "data": [
        "sql", "mysql", "postgresql", "sqlite", "oracle", "ms sql", "sql server",
        "nosql", "mongodb", "cassandra", "redis", "elasticsearch",
        "excel", "google sheets", "pivot tables", "vlookup", "power query",
        "tableau", "power bi", "looker", "qlikview", "qliksense", "metabase",
        "data analysis", "data analytics", "analytics", "data visualization",
        "data wrangling", "data cleaning", "etl", "data pipeline",
        "data engineering", "data modeling", "data warehousing",
        "big data", "hadoop", "spark", "kafka", "airflow", "dbt",
        "snowflake", "redshift", "bigquery", "databricks", "hive",
        "pandas", "numpy", "matplotlib", "seaborn", "plotly",
        "statistics", "statistical analysis", "hypothesis testing",
        "regression", "time series", "forecasting", "a/b testing",
    ],
    "ml_ai": [
        "machine learning", "deep learning", "neural networks",
        "natural language processing", "nlp", "computer vision",
        "reinforcement learning", "supervised learning", "unsupervised learning",
        "classification", "clustering", "recommendation systems",
        "feature engineering", "model deployment", "mlops",
        "scikit-learn", "sklearn", "tensorflow", "keras", "pytorch",
        "xgboost", "lightgbm", "catboost", "hugging face", "transformers",
        "bert", "gpt", "llm", "generative ai", "langchain", "rag",
        "opencv", "yolo", "object detection",
    ],
    "cloud_devops": [
        "aws", "azure", "gcp", "google cloud", "cloud computing",
        "docker", "kubernetes", "terraform", "ansible", "chef", "puppet",
        "ci/cd", "jenkins", "github actions", "gitlab ci", "circleci",
        "linux", "unix", "nginx", "apache", "microservices",
        "serverless", "lambda", "azure functions", "devops",
        "infrastructure as code", "prometheus", "grafana",
    ],
    "databases": [
        "database design", "database administration", "dba",
        "stored procedures", "indexing", "query optimization",
        "data migration",
    ],
    "finance": [
        "finance", "financial analysis", "financial modeling", "financial reporting",
        "accounting", "bookkeeping", "accounts payable", "accounts receivable",
        "general ledger", "reconciliation",
        "budgeting", "variance analysis", "cost analysis",
        "valuation", "dcf", "discounted cash flow", "equity research",
        "investment banking", "private equity", "venture capital",
        "portfolio management", "risk management", "credit analysis",
        "audit", "taxation", "tax", "gst", "tds",
        "ifrs", "gaap", "us gaap", "ind as",
        "tally", "quickbooks", "sap fico", "oracle financials",
        "ms dynamics", "xero", "zoho books",
        "balance sheet", "income statement", "cash flow statement",
        "working capital", "capex", "ebitda", "irr", "npv",
        "mergers and acquisitions", "m&a",
    ],
    "marketing": [
        "marketing", "digital marketing", "performance marketing",
        "seo", "sem", "search engine optimization", "search engine marketing",
        "google ads", "google adwords", "facebook ads", "meta ads",
        "instagram ads", "linkedin ads", "programmatic advertising",
        "content marketing", "content strategy", "content creation",
        "social media marketing", "social media management",
        "email marketing", "marketing automation", "crm",
        "hubspot", "salesforce", "marketo", "mailchimp", "klaviyo",
        "branding", "brand management", "brand strategy",
        "market research", "consumer insights", "competitive analysis",
        "product marketing", "go to market", "gtm strategy",
        "growth hacking", "growth marketing", "retention marketing",
        "affiliate marketing", "influencer marketing",
        "ahrefs", "semrush", "moz", "google analytics",
        "google tag manager", "mixpanel", "amplitude", "clevertap",
        "conversion rate optimization", "cro",
        "copywriting", "ad copywriting",
    ],
    "sales": [
        "sales", "b2b sales", "b2c sales", "inside sales", "field sales",
        "business development", "lead generation", "prospecting",
        "cold calling", "cold emailing", "outbound sales",
        "account management", "key account management", "kam",
        "client relationship management", "customer success",
        "revenue generation", "deal closing", "negotiation",
        "pipeline management", "crm management",
        "salesforce crm", "zoho crm", "hubspot crm",
    ],
    "product_project": [
        "product management", "product roadmap", "product strategy",
        "agile", "scrum", "kanban", "sprint planning", "backlog grooming",
        "jira", "confluence", "trello", "asana", "notion",
        "project management", "pmp", "prince2",
        "stakeholder management", "requirement gathering",
        "user stories", "mvp", "product launch",
        "ux", "ui", "user experience", "user interface",
        "wireframing", "prototyping", "figma", "sketch", "zeplin",
        "usability testing", "user research",
    ],
    "hr": [
        "human resources", "talent acquisition", "recruitment",
        "sourcing", "talent management", "performance management",
        "learning and development", "training",
        "compensation and benefits", "payroll", "hris",
        "employee engagement", "employee relations",
        "organizational development", "change management",
        "hr analytics", "workforce planning",
        "workday", "successfactors", "bamboohr", "darwinbox",
        "linkedin recruiter",
    ],
    "operations": [
        "operations management", "supply chain", "supply chain management",
        "procurement", "vendor management",
        "inventory management", "warehouse management",
        "logistics", "last mile delivery", "fulfillment",
        "lean", "six sigma", "kaizen", "process improvement",
        "quality assurance", "quality control",
        "erp", "sap", "oracle erp",
    ],
    "soft_skills": [
        "communication", "written communication", "verbal communication",
        "presentation", "public speaking", "leadership", "team leadership",
        "problem solving", "critical thinking", "decision making",
        "time management", "project coordination",
        "client communication", "stakeholder communication",
    ],
    "research": [
        "research", "literature review", "research methodology",
        "quantitative research", "qualitative research",
        "survey design", "data collection", "report writing",
        "academic writing", "peer review",
    ],
    "legal": [
        "legal research", "contract drafting", "contract review",
        "compliance", "regulatory compliance", "gdpr", "data privacy",
        "intellectual property", "corporate law", "litigation",
    ],
    "healthcare": [
        "clinical research", "clinical trials", "pharmacovigilance",
        "medical writing", "regulatory affairs", "gcp", "gmp",
        "healthcare management", "hospital management", "ehr", "emr",
    ],
    "design": [
        "graphic design", "visual design", "ui design", "ux design",
        "adobe photoshop", "photoshop", "illustrator", "adobe illustrator",
        "indesign", "after effects", "premiere pro",
        "canva", "figma", "sketch", "invision",
        "video editing", "motion graphics", "3d modeling",
        "autocad", "solidworks", "catia",
    ],
}

# ---------------------------------------------------------------------------
# Pre-computed flat structures derived from SKILL_LIBRARY
# ---------------------------------------------------------------------------

_ALL_SKILLS: set[str] = {
    skill.lower().strip()
    for skills in SKILL_LIBRARY.values()
    for skill in skills
}

_SKILLS_BY_LENGTH: list[str] = sorted(_ALL_SKILLS, key=len, reverse=True)

_SKILL_PATTERNS: dict[str, re.Pattern] = {
    skill: re.compile(
        r"(?<![a-z0-9\-\+\#\.])" + re.escape(skill) + r"(?![a-z0-9\-\+\#\.])",
        re.IGNORECASE,
    )
    for skill in _SKILLS_BY_LENGTH
}


# ===========================================================================
# EDUCATION MAP
# ===========================================================================

EDUCATION_DEGREES: list[tuple[list[str], int]] = [
    (["phd", "ph.d", "doctorate", "doctor of philosophy"], 100),
    (["m.tech", "mtech", "master of technology", "master of engineering"], 90),
    (["mba", "master of business administration", "masters in business"], 88),
    (["master", "masters", "m.s.", "m.sc", "msc", "master of science",
      "master of arts", "m.a."], 85),
    (["b.tech", "btech", "b.e.", "bachelor of technology",
      "bachelor of engineering"], 75),
    (["bachelor", "bachelors", "b.sc", "bsc", "bachelor of science",
      "b.com", "bcom", "bachelor of commerce", "b.a.", "bba",
      "bachelor of arts", "bachelor of business"], 70),
    (["diploma", "associate degree", "associate"], 50),
    (["12th", "hsc", "higher secondary", "intermediate", "senior secondary"], 35),
    (["10th", "ssc", "secondary school", "matric"], 20),
]


# ===========================================================================
# EXPERIENCE EXTRACTION
# ===========================================================================

_EXP_PATTERNS: list[re.Pattern] = [
    re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*years?\s+(?:of\s+)?(?:experience|exp\b)", re.I),
    re.compile(r"experience\s*(?:of\s*)?(\d+(?:\.\d+)?)\s*\+?\s*years?", re.I),
    re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*yrs?\s+(?:of\s+)?(?:experience|exp\b)", re.I),
    re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*years?\s+(?:of\s+)?(?:work|industry|relevant|professional)", re.I),
    re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*years?\s+exp", re.I),
]

# UPDATED: date-range patterns for AI-assisted experience calculation
_DATE_RANGE_PATTERN = re.compile(
    r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]+\d{4}|"
    r"\d{4})\s*(?:–|-|to)\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]+\d{4}|"
    r"\d{4}|present|current|till date|ongoing)",
    re.I,
)

_SKILLS_HEADERS: set[str] = {
    "skills", "technical skills", "core competencies",
    "technologies", "tech stack", "tools", "expertise",
    "competencies", "proficiencies", "key skills",
    "areas of expertise", "tools & technologies",
}

_OTHER_SECTIONS: set[str] = {
    "education", "experience", "work experience", "employment history",
    "projects", "certifications", "awards", "publications",
    "interests", "hobbies", "references", "summary", "objective",
    "profile", "about", "achievements", "languages", "extracurricular",
    "volunteering", "training", "courses",
}

_SKILLS_HDR_RE = re.compile(
    r"^\s*(" + "|".join(re.escape(h) for h in _SKILLS_HEADERS) + r")\s*[:\-]?\s*$",
    re.IGNORECASE,
)
_OTHER_HDR_RE = re.compile(
    r"^\s*(" + "|".join(re.escape(h) for h in _OTHER_SECTIONS) + r")\s*[:\-]?\s*$",
    re.IGNORECASE,
)

_NOISE_RES: list[re.Pattern] = [
    re.compile(r"[\$₹€£¥]\s*[\d,]+(?:\.\d+)?(?:k|l|lpa|lakh)?", re.I),
    re.compile(r"\d[\d,]*\s*(?:k|lpa|lakh|lac|cr|crore)?(?:/month|/year|per month|p\.m\.?|p\.a\.?)", re.I),
    re.compile(r"(?<![a-zA-Z])\b\d{4}\b"),
    re.compile(r"(?<![a-zA-Z\.])\b\d{1,3}\b(?!\s*[a-zA-Z%])"),
    re.compile(
        r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]+\d{4}\b",
        re.I,
    ),
    re.compile(r"stipend|salary|ctc|lpa|per annum|per month", re.I),
    re.compile(r"[^a-zA-Z0-9\s\.\-\+\#/&]"),
]


# ===========================================================================
# TEXT CLEANING
# ===========================================================================

def clean_text(text: str) -> str:
    for pat in _NOISE_RES:
        text = pat.sub(" ", text)
    return re.sub(r"\s+", " ", text).strip()


# ===========================================================================
# TEXT EXTRACTION
# ===========================================================================

def _read_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                pt = page.extract_text()
                if pt:
                    text += pt + "\n"
    except Exception:
        pass
    return text


def _read_docx(file_bytes: bytes) -> str:
    text = ""
    try:
        d = docx.Document(io.BytesIO(file_bytes))
        for para in d.paragraphs:
            text += para.text + "\n"
    except Exception:
        pass
    return text


def _read_txt(file_bytes: bytes) -> str:
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception:
        return ""


def extract_text(uploaded_file) -> str:
    """Extract raw text from an uploaded file (PDF / DOCX / TXT)."""
    uploaded_file.seek(0)
    data = uploaded_file.read()
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        return _read_pdf(data)
    if name.endswith(".docx"):
        return _read_docx(data)
    if name.endswith(".txt"):
        return _read_txt(data)
    return ""


# ===========================================================================
# CANDIDATE NAME EXTRACTION  (improved)
# Scans first 10-15 non-empty lines for a short, purely alphabetic,
# capitalized name line. Falls back to cleaned file stem.
# ===========================================================================

_NAME_EXCLUDE_WORDS: set[str] = {
    "academic", "education", "profile", "resume", "experience",
    "curriculum", "vitae", "objective", "summary", "contact",
    "phone", "email", "address", "linkedin", "github", "mobile",
    "www", "http", "about", "skills", "overview", "introduction",
}

_NAME_NOISE_CHARS_RE = re.compile(r"[^a-zA-Z\s]")


def extract_candidate_name(raw_text: str, file_stem: str = "") -> str:
    """
    Improved candidate name extraction.

    Strategy:
    1. Scan first 10-15 non-empty lines.
    2. A valid name line must:
       - Contain only alphabets and spaces (no digits, no special chars)
       - Have 2-4 words
       - NOT contain any excluded section/keyword words
       - Have all words capitalized (title-case preferred)
    3. Among valid candidates, prefer lines where all words start with uppercase.
    4. Fallback: clean version of the file stem.
    """
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    candidates = lines[:15]

    title_case_matches: list[str] = []
    any_matches: list[str] = []

    for line in candidates:
        # Strip leading bullet / special chars
        line = re.sub(r"^[\W_]+", "", line).strip()
        if not line:
            continue

        # Must contain only alphabets and spaces
        if _NAME_NOISE_CHARS_RE.search(line):
            continue

        # Must not contain excluded keywords
        lower_line = line.lower()
        if any(excl in lower_line for excl in _NAME_EXCLUDE_WORDS):
            continue

        words = line.split()

        # Must be 2–4 words
        if not (2 <= len(words) <= 4):
            continue

        # All words must be alphabetic only
        if not all(w.isalpha() for w in words):
            continue

        # Prefer title-cased (all words start with uppercase)
        if all(w[0].isupper() for w in words):
            title_case_matches.append(" ".join(words))
        else:
            any_matches.append(" ".join(words))

    if title_case_matches:
        return title_case_matches[0]
    if any_matches:
        return any_matches[0].title()

    # Fallback: humanise the file stem
    if file_stem:
        name = re.sub(r"[_\-\.]+", " ", file_stem).strip()
        # Remove file extension remnants
        name = re.sub(r"\s+(pdf|docx|txt)$", "", name, flags=re.I).strip()
        return name.title()

    return "Unknown Candidate"


# ===========================================================================
# SKILL EXTRACTION  (library-only, no free-form tokens)
# ===========================================================================

def find_skills_in_text(text: str) -> set[str]:
    tl = text.lower()
    found: set[str] = set()
    for skill in _SKILLS_BY_LENGTH:
        if _SKILL_PATTERNS[skill].search(tl):
            found.add(skill)
    return found


def _skills_section_bounds(lines: list[str]) -> tuple[int, int]:
    start = -1
    for i, line in enumerate(lines):
        if _SKILLS_HDR_RE.match(line.strip()):
            start = i
            break
    if start == -1:
        return -1, -1

    end = len(lines)
    for i in range(start + 1, len(lines)):
        stripped = lines[i].strip()
        if stripped and _OTHER_HDR_RE.match(stripped):
            end = i
            break
    return start, end


def extract_skills_weighted(raw_text: str, target_skills: set[str]) -> dict[str, float]:
    lines = raw_text.split("\n")
    s, e = _skills_section_bounds(lines)

    sec_lower = " ".join(lines[s:e]).lower() if s != -1 else ""
    full_lower = raw_text.lower()

    matched: dict[str, float] = {}
    for skill in target_skills:
        pat = _SKILL_PATTERNS[skill]
        if sec_lower and pat.search(sec_lower):
            matched[skill] = 1.0
        elif pat.search(full_lower):
            matched[skill] = 0.6
    return matched


# ===========================================================================
# JD FEATURE EXTRACTION  (called once per JD)
# ===========================================================================

def extract_jd_features(raw_text: str) -> dict:
    """
    Parse a JD once. Returns:
      skills         – set of SKILL_LIBRARY skills present in the JD
      required_exp   – years of experience required (0 if unspecified)
    """
    cleaned = clean_text(raw_text)
    skills = find_skills_in_text(cleaned)
    required_exp = _max_years(raw_text)
    return {"skills": skills, "required_exp": required_exp}


def _max_years(text: str) -> float:
    tl = text.lower()
    hits: list[float] = []
    for pat in _EXP_PATTERNS:
        for m in pat.findall(tl):
            try:
                hits.append(float(m))
            except ValueError:
                pass
    return max(hits) if hits else 0.0


# ===========================================================================
# RESUME PARSING
# ===========================================================================

def parse_resume(raw_text: str, file_stem: str = "") -> dict:
    """
    Parse a resume into structured fields.
    Returns: cleaned_text, raw_text, candidate_name, skills, experience_years,
             experience_display (human-readable), education_display (highest degree).

    # UPDATED: experience_display and education_display are now extracted
    # via Claude for more accurate, human-readable output in the Word report.
    """
    cleaned = clean_text(raw_text)
    base_years = _max_years(raw_text)

    # UPDATED: Use Claude to extract structured experience + education labels
    ai_meta = _extract_meta_with_claude(raw_text, base_years)

    return {
        "raw_text": raw_text,
        "cleaned_text": cleaned,
        "candidate_name": extract_candidate_name(raw_text, file_stem),
        "skills": find_skills_in_text(cleaned),
        "experience_years": ai_meta.get("experience_years", base_years),
        "experience_display": ai_meta.get("experience_display", _years_to_display(base_years)),
        "education_display": ai_meta.get("education_display", "Not specified"),
    }


# NEW: Convert decimal years to "X years Y months" string
def _years_to_display(years: float) -> str:
    """Convert 2.5 → '2 years 6 months', 3.0 → '3 years'."""
    if years <= 0:
        return "Not specified"
    total_months = round(years * 12)
    y = total_months // 12
    m = total_months % 12
    parts = []
    if y:
        parts.append(f"{y} year{'s' if y != 1 else ''}")
    if m:
        parts.append(f"{m} month{'s' if m != 1 else ''}")
    return " ".join(parts) if parts else "Not specified"


# NEW: Claude-powered extraction of experience (normalised) + highest education
_META_PROMPT_TEMPLATE = """\
You are a precise resume parser. Extract ONLY the following from the resume text below.

TASK 1 — TOTAL EXPERIENCE:
- Sum all work experience durations (exclude internships under 3 months, education periods)
- Handle date ranges like "Jan 2020 – Mar 2023", "2.5 years", "3+ years"
- Convert everything to decimal years (e.g. 2 years 6 months = 2.5)
- If resume mentions an explicit total (e.g. "5 years of experience"), trust it
- If no experience found, output 0

TASK 2 — HIGHEST EDUCATION:
- Identify the single highest qualification
- Format as: "<Degree> in <Field>" e.g. "B.Tech in Computer Science", "MBA – Marketing", "MSc Data Science", "PhD in Physics"
- If field is unknown, just output degree name e.g. "Bachelor of Commerce"
- If no education found, output "Not specified"

Hint: base rule-based experience estimate = {base_years} years (use as sanity check only)

RESUME:
{resume_snippet}

Respond ONLY in this exact format (no extra text):
EXPERIENCE_YEARS: <decimal number>
EXPERIENCE_DISPLAY: <X years Y months>
EDUCATION: <highest qualification string>
"""


def _extract_meta_with_claude(raw_text: str, base_years: float) -> dict:
    """
    # NEW: Use Claude to extract normalised experience + highest education.
    Sends only the first 3000 chars of resume (sufficient for dates + education).
    Falls back gracefully to rule-based values on any error.
    """
    fallback = {
        "experience_years": base_years,
        "experience_display": _years_to_display(base_years),
        "education_display": _rule_based_education(raw_text),
    }

    snippet = raw_text[:3000].strip()
    if not snippet:
        return fallback

    prompt = _META_PROMPT_TEMPLATE.format(
        base_years=round(base_years, 1),
        resume_snippet=snippet,
    )

    try:
        client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=120,
            messages=[{"role": "user", "content": prompt}],
        )
        return _parse_meta_response(msg.content[0].text.strip(), base_years)
    except Exception:
        return fallback


def _parse_meta_response(raw: str, base_years: float) -> dict:
    """
    # NEW: Parse structured Claude response for experience + education.
    Expected lines:
      EXPERIENCE_YEARS: 3.5
      EXPERIENCE_DISPLAY: 3 years 6 months
      EDUCATION: B.Tech in Computer Science
    """
    result: dict = {}
    for line in raw.splitlines():
        line = line.strip()
        if line.upper().startswith("EXPERIENCE_YEARS:"):
            val = line.split(":", 1)[1].strip()
            try:
                result["experience_years"] = float(re.findall(r"\d+(?:\.\d+)?", val)[0])
            except (IndexError, ValueError):
                result["experience_years"] = base_years
        elif line.upper().startswith("EXPERIENCE_DISPLAY:"):
            val = line.split(":", 1)[1].strip()
            result["experience_display"] = val if val and val.lower() != "none" else _years_to_display(base_years)
        elif line.upper().startswith("EDUCATION:"):
            val = line.split(":", 1)[1].strip()
            result["education_display"] = val if val and val.lower() not in ("none", "not specified", "") else "Not specified"

    # Fill any missing keys with fallbacks
    if "experience_years" not in result:
        result["experience_years"] = base_years
    if "experience_display" not in result:
        result["experience_display"] = _years_to_display(result["experience_years"])
    if "education_display" not in result:
        result["education_display"] = "Not specified"

    return result


def _rule_based_education(raw_text: str) -> str:
    """
    # NEW: Rule-based fallback for education label when Claude is unavailable.
    Returns the highest matched degree label.
    """
    tl = raw_text.lower()
    labels = [
        (["phd", "ph.d", "doctorate"], "PhD"),
        (["m.tech", "mtech", "master of technology", "master of engineering"], "M.Tech"),
        (["mba", "master of business administration"], "MBA"),
        (["msc", "m.sc", "master of science"], "MSc"),
        (["ma", "m.a.", "master of arts"], "MA"),
        (["masters", "master"], "Master's"),
        (["b.tech", "btech", "bachelor of technology", "b.e.", "bachelor of engineering"], "B.Tech"),
        (["bsc", "b.sc", "bachelor of science"], "BSc"),
        (["bcom", "b.com", "bachelor of commerce"], "B.Com"),
        (["bba", "bachelor of business"], "BBA"),
        (["ba", "b.a.", "bachelor of arts"], "BA"),
        (["bachelor", "bachelors"], "Bachelor's"),
        (["diploma"], "Diploma"),
    ]
    for keywords, label in labels:
        for kw in keywords:
            if re.search(r"(?<![a-z])" + re.escape(kw) + r"(?![a-z])", tl):
                return label
    return "Not specified"


# ===========================================================================
# SCORING  (base rule-based — unchanged)
# ===========================================================================

def compute_skill_score(
    raw_text: str, jd_skills: set[str]
) -> tuple[float, list[str], list[str]]:
    """
    Skill Score = sum(section-aware weights for matched JD skills)
                  ─────────────────────────────────────────────── × 100
                         total number of JD skills

    Returns (score_0_100, matched_list, missing_list).
    """
    if not jd_skills:
        return 0.0, [], []

    weighted = extract_skills_weighted(raw_text, jd_skills)
    raw_score = sum(weighted.values()) / len(jd_skills)
    score = round(min(raw_score * 100, 100.0), 1)
    matched = sorted(weighted.keys())
    missing = sorted(jd_skills - weighted.keys())
    return score, matched, missing


def compute_education_score(raw_text: str) -> float:
    tl = raw_text.lower()
    for keywords, pts in EDUCATION_DEGREES:
        for kw in keywords:
            pat = r"(?<![a-z])" + re.escape(kw) + r"(?![a-z])"
            if re.search(pat, tl):
                return float(pts)
    return 0.0


def compute_experience_score(resume_years: float, required_years: float) -> float:
    if required_years <= 0:
        return 100.0
    if resume_years >= required_years:
        return 100.0
    return round((resume_years / required_years) * 100, 1)


def score_resume_against_jd(
    resume: dict, jd_features: dict, weights: dict
) -> dict:
    """
    Composite weighted base score for a single (resume, JD) pair.
    Returns raw base scores + matched/missing skill lists.
    Claude refinement is applied separately to keep this function pure.
    """
    skill_score, matched, missing = compute_skill_score(
        resume["raw_text"], jd_features["skills"]
    )
    edu_score = compute_education_score(resume["raw_text"])
    exp_score = compute_experience_score(
        resume["experience_years"], jd_features["required_exp"]
    )

    total = (
        skill_score * (weights["skills"]     / 100)
        + edu_score  * (weights["education"]  / 100)
        + exp_score  * (weights["experience"] / 100)
    )

    return {
        "total":            round(total, 1),
        "skill_score":      round(skill_score, 1),
        "education_score":  round(edu_score, 1),
        "experience_score": round(exp_score, 1),
        "matched_skills":   matched,
        "missing_skills":   missing,
    }


# ===========================================================================
# CLUSTERING  — one resume → one JD, no duplicates
# ===========================================================================

def cluster_resumes_to_jds(
    resumes: list[dict],
    jd_list: list[dict],
    weights: dict,
    top_n: int,
) -> dict[str, list[dict]]:
    """
    Greedy assignment:
    1. Score every (resume, JD) pair.
    2. Sort all pairs descending by score.
    3. Walk the list: assign each resume to its best available JD.
       Each resume is assigned exactly once.
       Each JD accepts at most `top_n` candidates.

    Returns {jd_name: [{"name": str, "candidate_name": str, "score_data": dict}, ...]}
    sorted desc.
    """
    matrix: list[tuple[int, int, float, dict]] = []
    for r_idx, resume in enumerate(resumes):
        for j_idx, jd in enumerate(jd_list):
            sd = score_resume_against_jd(resume, jd["features"], weights)
            matrix.append((r_idx, j_idx, sd["total"], sd))

    matrix.sort(key=lambda x: x[2], reverse=True)

    assigned: set[int] = set()
    slots: dict[str, int] = defaultdict(int)
    result: dict[str, list[dict]] = defaultdict(list)

    for r_idx, j_idx, _, sd in matrix:
        if r_idx in assigned:
            continue
        jd_name = jd_list[j_idx]["name"]
        if slots[jd_name] >= top_n:
            continue
        result[jd_name].append({
            "name":               resumes[r_idx]["name"],
            "candidate_name":     resumes[r_idx].get("candidate_name", resumes[r_idx]["name"]),
            "experience_display": resumes[r_idx].get("experience_display", "Not specified"),  # UPDATED
            "education_display":  resumes[r_idx].get("education_display", "Not specified"),   # UPDATED
            "score_data":         sd,
        })
        assigned.add(r_idx)
        slots[jd_name] += 1
        if len(assigned) == len(resumes):
            break

    for name in result:
        result[name].sort(key=lambda x: x["score_data"]["total"], reverse=True)

    return dict(result)


# ===========================================================================
# CLAUDE API — SCORE REFINEMENT + CATEGORISED INSIGHTS
# Sends ONLY structured skill data (no resume / JD text). Token-optimised.
# One API call per shortlisted candidate.
# Final Score = 70% base + 30% Claude-refined.
# ===========================================================================

# UPDATED: Improved Claude prompt — recruiter-style, no category labels,
# penalises missing critical skills, rewards depth over keyword matching.
# Output is structured for deterministic parsing into bullet lists.
_REFINE_PROMPT_TEMPLATE = """\
You are a senior technical recruiter. Evaluate this candidate's fit for the role.

JD Required Skills: {jd_skills}
Candidate Matched Skills: {matched_skills}
Candidate Missing Skills: {missing_skills}
Candidate Experience: {experience_display}
Candidate Education: {education_display}

Base Scores (rule-based, 0-100): Skills={skill_score} Education={education_score} Experience={experience_score}

INSTRUCTIONS:
1. Slightly adjust scores based on quality signals, not just keyword presence
2. Penalise missing critical/core skills (first 5 in JD skills list are most critical)
3. Reward demonstrated depth (multiple related skills > isolated keywords)
4. Keep score adjustments within ±15 points of base
5. Write recruiter-style strengths and gaps — short, sharp, non-generic

STRENGTHS — write exactly 3 bullet points:
- Focus on what the candidate genuinely brings
- Be specific to this JD, not generic praise

GAPS — write exactly 3 bullet points:
- Focus on what's actually missing or weak for this role
- No softening language — direct recruiter voice

Output ONLY in this exact format (no extra text, no headers, no categories):
Refined Scores: Skills: XX Education: XX Experience: XX Final: XX
Strengths:
- <point 1>
- <point 2>
- <point 3>
Gaps:
- <point 1>
- <point 2>
- <point 3>\
"""


def refine_scores_with_claude(
    jd_skills: list[str],
    matched_skills: list[str],
    missing_skills: list[str],
    base_skill_score: float,
    base_education_score: float,
    base_experience_score: float,
    base_total: float,
    weights: dict,
    experience_display: str = "Not specified",  # UPDATED: pass human-readable labels
    education_display: str = "Not specified",   # UPDATED
) -> dict:
    """
    Send compact structured data to Claude. Returns:
      refined_skill_score, refined_education_score, refined_experience_score,
      refined_total, blended_total, strengths (list), gaps (list).

    # UPDATED: prompt now includes experience_display + education_display for
    # richer context. Insights are now plain bullet lists (no category labels).
    Blend formula: blended_total = 0.70 × base_total + 0.30 × claude_total

    Falls back gracefully on any API error.
    """
    fallback = _build_fallback(
        base_skill_score, base_education_score, base_experience_score, base_total
    )

    prompt = _REFINE_PROMPT_TEMPLATE.format(
        jd_skills=jd_skills[:20],
        matched_skills=matched_skills[:15],
        missing_skills=missing_skills[:15],
        skill_score=round(base_skill_score),
        education_score=round(base_education_score),
        experience_score=round(base_experience_score),
        experience_display=experience_display,
        education_display=education_display,
    )

    try:
        client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        message = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=400,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = message.content[0].text.strip()
        return _parse_refinement(
            raw, base_skill_score, base_education_score,
            base_experience_score, base_total, weights,
        )
    except Exception:
        return fallback


def _build_fallback(
    skill: float, edu: float, exp: float, total: float
) -> dict:
    # UPDATED: strengths/gaps are now plain lists (no category keys)
    return {
        "refined_skill_score":      skill,
        "refined_education_score":  edu,
        "refined_experience_score": exp,
        "refined_total":            total,
        "blended_total":            total,
        "strengths": [
            "Matched key technical skills for the role",
            "Experience level aligns with role requirements",
            "Educational background supports the position",
        ],
        "gaps": [
            "Some required technical skills not verified",
            "Additional domain experience may be needed",
            "Seniority fit requires further assessment",
        ],
    }


def _parse_refinement(
    raw: str,
    base_skill: float,
    base_edu: float,
    base_exp: float,
    base_total: float,
    weights: dict,
) -> dict:
    """
    # UPDATED: Parse Claude's plain-text output into refined scores + bullet lists.
    New format expects:
      Refined Scores: Skills: XX Education: XX Experience: XX Final: XX
      Strengths:
      - <point>
      - <point>
      - <point>
      Gaps:
      - <point>
      - <point>
      - <point>

    Old category-prefixed format (Skills:/Education:/etc.) is no longer expected.
    Falls back to base values on any parse failure.
    """
    refined_skill = base_skill
    refined_edu   = base_edu
    refined_exp   = base_exp
    refined_total = base_total

    strengths: list[str] = []
    gaps:      list[str] = []
    current_section: str | None = None  # "strengths" or "gaps"

    for line in raw.splitlines():
        line = line.strip()
        if not line:
            continue
        lower = line.lower()

        # ── Refined scores line ──────────────────────────────────────────────
        if lower.startswith("refined scores"):
            nums = re.findall(r"(\d+(?:\.\d+)?)", line)
            if len(nums) >= 4:
                refined_skill = _clamp(float(nums[0]))
                refined_edu   = _clamp(float(nums[1]))
                refined_exp   = _clamp(float(nums[2]))
                refined_total = _clamp(float(nums[3]))
            elif len(nums) == 3:
                refined_skill = _clamp(float(nums[0]))
                refined_edu   = _clamp(float(nums[1]))
                refined_exp   = _clamp(float(nums[2]))
                refined_total = (
                    refined_skill * (weights["skills"]     / 100)
                    + refined_edu   * (weights["education"]  / 100)
                    + refined_exp   * (weights["experience"] / 100)
                )
            continue

        # ── Section headers ──────────────────────────────────────────────────
        if lower.startswith("strengths"):
            current_section = "strengths"
            continue
        if lower.startswith("gaps"):
            current_section = "gaps"
            continue

        # ── Plain bullet items ───────────────────────────────────────────────
        if line.startswith(("-", "*", "•")) and current_section is not None:
            text = line.lstrip("-*• ").strip()
            if not text:
                continue
            # UPDATED: Strip any residual "Category: " prefix if present
            text = re.sub(r"^(Skills|Education|Experience|Overall)\s*[:\-]\s*", "", text, flags=re.I).strip()
            if current_section == "strengths" and len(strengths) < 3:
                strengths.append(text)
            elif current_section == "gaps" and len(gaps) < 3:
                gaps.append(text)

    # Fill missing bullets with sensible defaults
    _default_strengths = [
        "Matched key technical skills for the role",
        "Experience level aligns with role requirements",
        "Educational background supports the position",
    ]
    _default_gaps = [
        "Some required technical skills not verified",
        "Additional domain experience may be needed",
        "Seniority fit requires further assessment",
    ]
    while len(strengths) < 3:
        strengths.append(_default_strengths[len(strengths)])
    while len(gaps) < 3:
        gaps.append(_default_gaps[len(gaps)])

    # Blend: 70% base + 30% Claude
    blended_total = round(0.70 * base_total + 0.30 * refined_total, 1)

    return {
        "refined_skill_score":      round(refined_skill, 1),
        "refined_education_score":  round(refined_edu, 1),
        "refined_experience_score": round(refined_exp, 1),
        "refined_total":            round(refined_total, 1),
        "blended_total":            blended_total,
        "strengths":                strengths,
        "gaps":                     gaps,
    }


def _clamp(v: float, lo: float = 0.0, hi: float = 100.0) -> float:
    return max(lo, min(hi, v))


# ===========================================================================
# WORD REPORT GENERATION
# ===========================================================================

def _rgb(run, r: int, g: int, b: int) -> None:
    run.font.color.rgb = RGBColor(r, g, b)


def _hr(doc: Document) -> None:
    """Thin grey horizontal rule between candidate blocks."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot)
    pPr.append(pBdr)


def _set_cell_bg(cell, hex_color: str) -> None:
    """Apply a solid background colour to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell) -> None:
    """Thin border on all four sides of a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "CCCCCC")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _score_table(doc: Document, refinement: dict, experience_display: str, education_display: str) -> None:
    """
    # UPDATED: Compact 2-column info table showing human-readable values.
    REMOVED percentage-based scores for Skills/Education/Experience.
    Now shows:
      Experience  | X years Y months
      Education   | Highest qualification
    Skills match is shown separately via _skills_table().
    """
    rows_data = [
        ("Experience",  experience_display),
        ("Education",   education_display),
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=2)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for cell, text in zip(hdr_cells, ("Metric", "Details")):
        _set_cell_bg(cell, "1F457C")
        _set_cell_border(cell)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, (label, value) in enumerate(rows_data, 1):
        row_cells = table.rows[i].cells
        bg_hex = "F9FAFB" if i % 2 == 0 else "FFFFFF"

        for cell in row_cells:
            _set_cell_bg(cell, bg_hex)
            _set_cell_border(cell)

        # Label cell
        lp = row_cells[0].paragraphs[0]
        lp.clear()
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)

        # Value cell
        vp = row_cells[1].paragraphs[0]
        vp.clear()
        vr = vp.add_run(value)
        vr.font.size = Pt(10)


def _skills_table(doc: Document, matched_skills: list[str], missing_skills: list[str]) -> None:
    """
    # NEW: Two-column skills table — Matched Skills | Missing Skills.
    REMOVED percentage match column entirely.
    Skills are listed as plain comma-separated text per cell for readability.
    """
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for cell, label in zip(hdr_cells, ("✔ Matched Skills", "✘ Missing Skills")):
        _set_cell_bg(cell, "1F457C")
        _set_cell_border(cell)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Content row
    data_cells = table.rows[1].cells
    _set_cell_bg(data_cells[0], "F0FFF4")
    _set_cell_bg(data_cells[1], "FFF8F0")
    for cell in data_cells:
        _set_cell_border(cell)

    matched_text = ", ".join(sorted(matched_skills)) if matched_skills else "None"
    missing_text = ", ".join(sorted(missing_skills)) if missing_skills else "None"

    mp = data_cells[0].paragraphs[0]
    mp.clear()
    mr = mp.add_run(matched_text)
    mr.font.size = Pt(9)

    gp = data_cells[1].paragraphs[0]
    gp.clear()
    gr = gp.add_run(missing_text)
    gr.font.size = Pt(9)


def _final_score_block(doc: Document, blended_total: float) -> None:
    """
    Bold, larger-font final match score highlight paragraph.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"FINAL MATCH SCORE:  {blended_total}%")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x37, 0x86, 0x3C)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _insights_table(doc: Document, refinement: dict) -> None:
    """
    # UPDATED: Clean 2-column Strengths vs Gaps table.
    REMOVED category labels (Skills/Education/Experience/Overall).
    REPLACED WITH: 3 plain recruiter-style bullet points per column.

    | Strengths          | Gaps               |
    |--------------------|---------------------|
    | • point 1          | • point 1           |
    | • point 2          | • point 2           |
    | • point 3          | • point 3           |
    """
    strengths = refinement["strengths"]  # list[str]
    gaps      = refinement["gaps"]       # list[str]

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for cell, label in zip(hdr_cells, ("Strengths", "Gaps")):
        _set_cell_bg(cell, "1F457C")
        _set_cell_border(cell)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Single content row with multi-line bullet text
    content_row = table.add_row()
    s_cell = content_row.cells[0]
    g_cell = content_row.cells[1]
    _set_cell_bg(s_cell, "F0FFF4")
    _set_cell_bg(g_cell, "FFF8F0")
    _set_cell_border(s_cell)
    _set_cell_border(g_cell)

    def _fill_bullets(cell, items: list[str]) -> None:
        first = True
        for item in items:
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.clear()
            run = p.add_run(f"• {item}")
            run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(3)

    _fill_bullets(s_cell, strengths[:3])
    _fill_bullets(g_cell, gaps[:3])


def generate_jd_report(
    jd_name: str,
    jd_skills: list[str],
    candidates: list[dict],
    weights: dict,
) -> bytes:
    """
    Build and return a Word (.docx) report for one JD as raw bytes.

    # UPDATED per-candidate layout:
      1. HEADER           — candidate name + file name
      2. PROFILE TABLE    — Experience (X yrs Y months) + Education (highest degree)
      3. SKILLS TABLE     — Matched / Missing skills (no percentage column)
      4. FINAL SCORE      — bold highlighted blended total
      5. INSIGHTS TABLE   — clean 2-col Strengths vs Gaps (plain bullets, no categories)
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.2)
        sec.right_margin  = Inches(1.2)

    # ── Report title ──────────────────────────────────────────────────────────
    title = doc.add_heading(f"Hiring Report \u2014 {jd_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        _rgb(title.runs[0], 0x1F, 0x45, 0x7C)
    doc.add_paragraph()

    if not candidates:
        p = doc.add_paragraph("No suitable candidates found for this role.")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        s = doc.add_paragraph()
        s.add_run("Total Candidates Selected: ").bold = True
        s.add_run(str(len(candidates)))
        doc.add_paragraph()

        for idx, cand in enumerate(candidates, 1):
            sd               = cand["score_data"]
            file_stem        = cand["name"]
            candidate_name   = cand.get("candidate_name", file_stem)
            exp_display      = cand.get("experience_display", "Not specified")   # UPDATED
            edu_display      = cand.get("education_display", "Not specified")    # UPDATED

            # ── Candidate heading ─────────────────────────────────────────────
            h = doc.add_heading(f"{idx}. {candidate_name}", level=2)
            if h.runs:
                _rgb(h.runs[0], 0x2E, 0x74, 0xB5)

            # File name sub-label
            fn_p = doc.add_paragraph()
            fn_r = fn_p.add_run(f"File: {file_stem}")
            fn_r.font.size = Pt(9)
            fn_r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            fn_p.paragraph_format.space_after = Pt(4)

            doc.add_paragraph()

            # ── Claude refinement (one API call per candidate) ─────────────────
            # UPDATED: now passes experience_display + education_display to Claude
            refinement = refine_scores_with_claude(
                jd_skills=jd_skills,
                matched_skills=sd["matched_skills"],
                missing_skills=sd["missing_skills"],
                base_skill_score=sd["skill_score"],
                base_education_score=sd["education_score"],
                base_experience_score=sd["experience_score"],
                base_total=sd["total"],
                weights=weights,
                experience_display=exp_display,
                education_display=edu_display,
            )

            # ── Profile table (Experience + Education, no percentages) ─────────
            tbl_hdr = doc.add_paragraph()
            tbl_hdr.add_run("Candidate Profile").bold = True
            tbl_hdr.paragraph_format.space_after = Pt(4)

            _score_table(doc, refinement, exp_display, edu_display)   # UPDATED signature
            doc.add_paragraph()

            # ── Skills table (matched vs missing, no % column) ────────────────
            sk_hdr = doc.add_paragraph()
            sk_hdr.add_run("Skills Match").bold = True
            sk_hdr.paragraph_format.space_after = Pt(4)

            _skills_table(doc, sd["matched_skills"], sd["missing_skills"])  # NEW
            doc.add_paragraph()

            # ── Final score highlight ─────────────────────────────────────────
            _final_score_block(doc, refinement["blended_total"])
            doc.add_paragraph()

            # ── Insights table (clean bullets — no "Candidate Insights" heading) ─
            _insights_table(doc, refinement)   # UPDATED: heading removed per spec

            _hr(doc)
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ===========================================================================
# STREAMLIT UI
# ===========================================================================

def main() -> None:
    st.set_page_config(
        page_title="ATS \u2014 Resume Screener",
        page_icon="\U0001f4cb",
        layout="centered",
    )

    st.markdown(
        """
        <style>
        /* ── Container ── */
        .block-container {
            max-width: 760px;
            padding-top: 2.5rem;
            padding-bottom: 3.5rem;
        }

        /* ── Sidebar ── */
        [data-testid="stSidebar"] {
            border-right: 1px solid #2a2d36;
        }
        [data-testid="stSidebar"] > div:first-child {
            padding-top: 2rem;
        }

        /* ── Sidebar section cards ── */
        .sb-card {
            background: #0f1115;
            border: 1px solid #2a2d36;
            border-radius: 10px;
            padding: 1rem 1rem 0.25rem 1rem;
            margin-bottom: 0.9rem;
        }
        .sb-label {
            font-size: 0.68rem;
            font-weight: 700;
            letter-spacing: 0.1em;
            text-transform: uppercase;
            color: #555d6e;
            margin-bottom: 0.6rem;
        }

        /* ── Upload panels ── */
        .up-panel {
            background: #1a1d24;
            border: 1px solid #2a2d36;
            border-radius: 10px;
            padding: 1.1rem 1.1rem 0.5rem 1.1rem;
            min-height: 170px;
        }
        .up-label {
            font-size: 0.75rem;
            font-weight: 700;
            letter-spacing: 0.07em;
            text-transform: uppercase;
            color: #6b7280;
            margin-bottom: 0.5rem;
        }
        .up-badge {
            display: inline-block;
            background: #0d2b1f;
            color: #4ade80;
            border: 1px solid #166534;
            border-radius: 999px;
            font-size: 0.75rem;
            font-weight: 600;
            padding: 0.15rem 0.7rem;
            margin-top: 0.35rem;
        }

        /* ── Spacing helpers ── */
        .gap-sm { margin-top: 0.9rem; }
        .gap-md { margin-top: 1.6rem; }

        /* ── Primary button ── */
        button[kind="primary"] {
            border-radius: 8px !important;
            font-weight: 600 !important;
            letter-spacing: 0.02em !important;
        }

        /* ── Download buttons ── */
        [data-testid="stDownloadButton"] button {
            border-radius: 8px !important;
            font-weight: 500 !important;
        }

        /* ── Headings ── */
        h1 { font-size: 1.65rem !important; font-weight: 700 !important; margin-bottom: 0.15rem !important; }
        h3 { font-weight: 600 !important; margin-top: 0 !important; }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # -----------------------------------------------------------------------
    # SIDEBAR
    # -----------------------------------------------------------------------
    with st.sidebar:
        st.markdown("## ⚙️ Configuration")
        st.markdown("<div class='gap-sm'></div>", unsafe_allow_html=True)

        st.markdown(
            "<div class='sb-card'><div class='sb-label'>Top Candidates</div>",
            unsafe_allow_html=True,
        )
        top_n = st.slider("Per JD", min_value=1, max_value=20, value=5)
        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown(
            "<div class='sb-card'><div class='sb-label'>Scoring Weights (%)</div>",
            unsafe_allow_html=True,
        )
        w_skills = st.slider("Skills Weight",     0, 100, 60, step=5)
        w_edu    = st.slider("Education Weight",  0, 100, 20, step=5)
        w_exp    = st.slider("Experience Weight", 0, 100, 20, step=5)
        st.markdown("</div>", unsafe_allow_html=True)

        total_w = w_skills + w_edu + w_exp
        if total_w != 100:
            st.error(f"Total is **{total_w}%** — weights must sum to 100%.")
        else:
            st.success("✅ Weights sum to 100%")

    # -----------------------------------------------------------------------
    # MAIN PAGE
    # -----------------------------------------------------------------------
    st.title("📄 ATS Resume Screener")
    st.markdown(
        "<p style='color:#6b7280; font-size:0.93rem; margin-top:0.05rem; margin-bottom:0;'>"
        "Upload job descriptions and resumes — get ranked hiring reports instantly."
        "</p>",
        unsafe_allow_html=True,
    )

    st.markdown("<div class='gap-md'></div>", unsafe_allow_html=True)

    col_jd, col_res = st.columns(2, gap="medium")

    with col_jd:
        st.markdown(
            "<div class='up-panel'><div class='up-label'>📋 Job Descriptions</div>",
            unsafe_allow_html=True,
        )
        jd_files = st.file_uploader(
            "Upload JDs",
            type=["pdf", "docx", "txt"],
            accept_multiple_files=True,
            key="jd_uploader",
            label_visibility="collapsed",
        )
        if jd_files:
            n = len(jd_files)
            st.markdown(
                f"<div class='up-badge'>✅ {n} JD{'s' if n != 1 else ''} uploaded</div>",
                unsafe_allow_html=True,
            )
        st.markdown("</div>", unsafe_allow_html=True)

    with col_res:
        st.markdown(
            "<div class='up-panel'><div class='up-label'>📄 Resumes</div>",
            unsafe_allow_html=True,
        )
        resume_files = st.file_uploader(
            "Upload Resumes",
            type=["pdf", "docx", "txt"],
            accept_multiple_files=True,
            key="resume_uploader",
            label_visibility="collapsed",
        )
        if resume_files:
            n = len(resume_files)
            st.markdown(
                f"<div class='up-badge'>✅ {n} Resume{'s' if n != 1 else ''} uploaded</div>",
                unsafe_allow_html=True,
            )
        st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<div class='gap-md'></div>", unsafe_allow_html=True)

    if st.button("📄 Generate Reports", type="primary", use_container_width=True):
        errors = []
        if not jd_files:
            errors.append("Upload at least one Job Description.")
        if not resume_files:
            errors.append("Upload at least one Resume.")
        if total_w != 100:
            errors.append("Scoring weights in the sidebar must sum to 100%.")
        for e in errors:
            st.error(e)
        if errors:
            return

        weights = {"skills": w_skills, "education": w_edu, "experience": w_exp}

        with st.spinner("Parsing job descriptions…"):
            jd_list = []
            for f in jd_files:
                raw = extract_text(f)
                if not raw.strip():
                    st.warning(f"⚠️ Could not read: {f.name}")
                    continue
                jd_list.append({
                    "name":     f.name.rsplit(".", 1)[0],
                    "features": extract_jd_features(raw),
                })

        if not jd_list:
            st.error("No valid JDs parsed.")
            return

        with st.spinner("Parsing resumes…"):
            resume_list = []
            for f in resume_files:
                raw = extract_text(f)
                if not raw.strip():
                    st.warning(f"⚠️ Could not read: {f.name}")
                    continue
                file_stem = f.name.rsplit(".", 1)[0]
                parsed = parse_resume(raw, file_stem)
                resume_list.append({"name": file_stem, **parsed})

        if not resume_list:
            st.error("No valid resumes parsed.")
            return

        with st.spinner("Scoring and assigning candidates…"):
            assignments = cluster_resumes_to_jds(resume_list, jd_list, weights, top_n)

        st.markdown("<div class='gap-md'></div>", unsafe_allow_html=True)
        st.markdown("### Results Summary")

        for jd in jd_list:
            candidates = assignments.get(jd["name"], [])
            if candidates:
                st.success(f"**{jd['name']}** → {len(candidates)} candidate(s) selected")
            else:
                st.warning(f"**{jd['name']}** → No suitable candidates found")

        any_dl = any(assignments.get(jd["name"]) for jd in jd_list)
        if any_dl:
            st.markdown("<div class='gap-md'></div>", unsafe_allow_html=True)
            st.markdown("### Download Reports")
            for jd in jd_list:
                candidates = assignments.get(jd["name"], [])
                if not candidates:
                    continue
                jd_skills = sorted(jd["features"]["skills"])
                with st.spinner(f"Generating insights for {jd['name']}\u2026"):
                    report = generate_jd_report(
                        jd["name"], jd_skills, candidates, weights
                    )
                safe = re.sub(r"[^\w\-_]", "_", jd["name"])
                st.download_button(
                    label=f"⬇ Download — {jd['name']}",
                    data=report,
                    file_name=f"{safe}_report.docx",
                    mime=(
                        "application/vnd.openxmlformats-officedocument"
                        ".wordprocessingml.document"
                    ),
                    use_container_width=True,
                )
        else:
            st.warning("No candidates were assigned to any JD.")


if __name__ == "__main__":
    main()
